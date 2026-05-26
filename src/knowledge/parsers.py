from __future__ import annotations

import hashlib
from pathlib import Path

from src.knowledge.schemas import ParsedDocument, ParsedPage

SUPPORTED_DOCUMENT_SUFFIXES = {".md", ".txt", ".pdf", ".docx"}


class UnsupportedDocumentTypeError(ValueError):
    pass


class DocumentParseError(RuntimeError):
    pass


def parse_document(path: str | Path, *, document_id: str) -> ParsedDocument:
    document_path = Path(path)
    suffix = document_path.suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_SUFFIXES))
        raise UnsupportedDocumentTypeError(
            f"Unsupported document type '{suffix}'. Supported: {supported}"
        )

    if suffix in {".md", ".txt"}:
        pages = [_parse_text_page(document_path)]
    elif suffix == ".pdf":
        pages = _parse_pdf_pages(document_path)
    else:
        pages = _parse_docx_pages(document_path)

    if not any(page.normalized_text() for page in pages):
        raise DocumentParseError(
            f"Parsed document '{document_path.name}' did not contain extractable text."
        )

    return ParsedDocument(
        document_id=document_id,
        filename=document_path.name,
        file_type=suffix,
        file_hash=file_sha256(document_path),
        source_path=str(document_path),
        pages=pages,
        metadata={"filename": document_path.name, "source_path": str(document_path)},
    )


def file_sha256(path: str | Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _parse_text_page(path: Path) -> ParsedPage:
    text = path.read_text(encoding="utf-8")
    return ParsedPage(
        page_number=None,
        text=text,
        section_title=_first_heading(text),
    )


def _parse_pdf_pages(path: Path) -> list[ParsedPage]:
    try:
        return _parse_pdf_pages_with_pymupdf(path)
    except ImportError:
        return _parse_pdf_pages_with_pypdf(path)


def _parse_pdf_pages_with_pymupdf(path: Path) -> list[ParsedPage]:
    try:
        import fitz
    except ImportError as exc:
        raise ImportError("PyMuPDF is not installed.") from exc

    pages: list[ParsedPage] = []
    try:
        with fitz.open(str(path)) as document:
            for index, page in enumerate(document, start=1):
                pages.append(
                    ParsedPage(
                        page_number=index,
                        text=page.get_text("text") or "",
                        section_title=None,
                    )
                )
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse PDF '{path.name}': {exc}") from exc
    return pages


def _parse_pdf_pages_with_pypdf(path: Path) -> list[ParsedPage]:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise DocumentParseError(
            "PDF parsing requires PyMuPDF or pypdf. Install project dependencies."
        ) from exc

    try:
        reader = PdfReader(str(path))
        return [
            ParsedPage(page_number=index, text=page.extract_text() or "", section_title=None)
            for index, page in enumerate(reader.pages, start=1)
        ]
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse PDF '{path.name}': {exc}") from exc


def _parse_docx_pages(path: Path) -> list[ParsedPage]:
    try:
        from docx import Document
    except Exception as exc:
        raise DocumentParseError(
            "DOCX parsing requires python-docx. Install project dependencies."
        ) from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX '{path.name}': {exc}") from exc
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n".join(lines)
    return [ParsedPage(page_number=None, text=text, section_title=_first_heading(text))]


def _first_heading(text: str) -> str | None:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip() or None
        if stripped and len(stripped.split()) <= 12:
            return stripped
    return None
